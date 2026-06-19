#!/usr/bin/env python3

from __future__ import annotations

import argparse
import datetime as dt
import html
import json
import os
import re
import subprocess
import sys
import urllib.request
from pathlib import Path


ROOT = Path(__file__).resolve().parent.parent
RAW_SOURCES = ROOT / "raw" / "sources"
WIKI = ROOT / "wiki"
WIKI_SOURCES = WIKI / "sources"
WIKI_CONCEPTS = WIKI / "concepts"

NOISE_CANDIDATE_TAGS = {
    "source",
    "local",
    "web",
    "paper",
    "article",
    "note",
    "dataset",
    "video",
    "sciverse",
    "academic",
    "rag",
    "lightrag",
    "graphrag",
    "rag-anything",
    "inbox",
    "active",
    "preview",
    "validation",
}
WIKI_DOMAINS = WIKI / "domains"
DEFAULT_WECHAT_CRAWLER = Path(
    "/Users/Min369/Documents/同步空间/Manju/AIProjects/MZ数据库/心之髓/wechat_crawler"
)

PLACEHOLDERS = {
    "为什么要收录这份资料？它与当前 wiki 的哪条主线相关？",
    "把原文、摘录、转写或清洗后的正文放在这里。",
    "可补充少量人工备注，但尽量不要破坏原始内容的可追溯性。",
}

BLOCKED_TEXT_MARKERS = (
    "环境异常",
    "当前环境异常",
    "完成验证后即可继续访问",
    "去验证",
    "轻点两下取消赞",
    "轻点两下取消在看",
)


def parse_frontmatter_and_body(path: Path) -> tuple[dict[str, str], str]:
    lines = path.read_text(encoding="utf-8").splitlines()
    if not lines or lines[0].strip() != "---":
        return {}, "\n".join(lines)

    data: dict[str, str] = {}
    end_index = 0
    for index, line in enumerate(lines[1:], start=1):
        stripped = line.strip()
        if stripped == "---":
            end_index = index
            break
        if ":" not in stripped:
            continue
        key, value = stripped.split(":", 1)
        data[key.strip()] = value.strip()

    body = "\n".join(lines[end_index + 1 :]) if end_index else "\n".join(lines)
    return data, body


def parse_list(value: str | None) -> list[str]:
    if not value:
        return []
    cleaned = value.strip()
    if cleaned.startswith("[") and cleaned.endswith("]"):
        cleaned = cleaned[1:-1]
    items: list[str] = []
    for item in cleaned.split(","):
        normalized = item.strip().strip("\"'")
        if normalized and normalized not in items:
            items.append(normalized)
    return items


def slug_from_raw_path(path: Path) -> str:
    stem = path.stem
    match = re.match(r"^\d{4}-\d{2}-\d{2}-(.+)$", stem)
    return match.group(1) if match else stem


def date_from_meta_or_path(meta: dict[str, str], path: Path) -> str:
    created = meta.get("created", "").strip()
    if re.match(r"^\d{4}-\d{2}-\d{2}$", created):
        return created
    match = re.match(r"^(\d{4}-\d{2}-\d{2})-", path.name)
    return match.group(1) if match else dt.date.today().isoformat()


def extract_sections(body: str) -> dict[str, str]:
    sections: dict[str, list[str]] = {}
    current = "_intro"
    sections[current] = []
    for line in body.splitlines():
        match = re.match(r"^##\s+(.+?)\s*$", line)
        if match:
            current = match.group(1).strip()
            sections.setdefault(current, [])
            continue
        sections.setdefault(current, []).append(line)
    return {key: "\n".join(value).strip() for key, value in sections.items()}


def remove_placeholders(text: str) -> str:
    cleaned_lines = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append(line)
            continue
        if stripped in PLACEHOLDERS:
            continue
        if stripped.startswith("学科归属："):
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def clean_sciverse_original(text: str) -> str:
    cleaned_lines: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            cleaned_lines.append("")
            continue
        bullet = re.sub(r"^\s*[-*]\s+", "", line)
        if bullet in {"Sciverse academic retrieval result.", "Sciverse academic retrieval result"}:
            continue
        if re.match(r"^(?:Title|Doc ID|Authors?|Auth|Year|Venue|DOI|Score|Retrieval query|Offset)\s*:", bullet, re.I):
            continue
        if bullet.lower().startswith("excerpt:"):
            excerpt = bullet.split(":", 1)[1].strip()
            if excerpt:
                cleaned_lines.append(excerpt)
            continue
        cleaned_lines.append(line)
    return "\n".join(cleaned_lines).strip()


def plain_text(markdown: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", markdown)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def split_long_text(text: str, limit: int = 260) -> list[str]:
    cleaned = plain_text(text)
    if not cleaned:
        return []
    pieces = re.split(r"(?<=[。！？.!?])\s+|---+", cleaned)
    chunks: list[str] = []
    current = ""
    for piece in pieces:
        part = piece.strip()
        if not part:
            continue
        if len(part) > limit:
            for start in range(0, len(part), limit):
                sub = part[start : start + limit].strip()
                if sub:
                    chunks.append(sub)
            continue
        if current and len(current) + len(part) + 1 > limit:
            chunks.append(current)
            current = part
        else:
            current = f"{current} {part}".strip()
    if current:
        chunks.append(current)
    return chunks


def first_paragraphs(text: str, limit: int = 5) -> list[str]:
    paragraphs = []
    for block in re.split(r"\n\s*\n", remove_placeholders(text)):
        cleaned = plain_text(block)
        if len(cleaned) > 360:
            for chunk in split_long_text(cleaned):
                if chunk and chunk not in paragraphs:
                    paragraphs.append(chunk)
                if len(paragraphs) >= limit:
                    return paragraphs
            continue
        if cleaned and cleaned not in paragraphs:
            paragraphs.append(cleaned)
        if len(paragraphs) >= limit:
            break
    return paragraphs


def truncate(text: str, limit: int = 180) -> str:
    cleaned = plain_text(text)
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[: limit - 1].rstrip() + "…"


def is_blocked_or_boilerplate(text: str) -> bool:
    cleaned = plain_text(text)
    if not cleaned:
        return False
    marker_count = sum(1 for marker in BLOCKED_TEXT_MARKERS if marker in cleaned)
    if marker_count >= 2:
        return True
    if "环境异常" in cleaned and len(cleaned) < 500:
        return True
    return False


def markdown_list(items: list[str], fallback: str) -> list[str]:
    if not items:
        return [f"- {fallback}"]
    return [f"- {item}" for item in items]


def html_to_text(raw_html: str) -> str:
    text = re.sub(r"(?is)<(script|style|noscript).*?</\1>", " ", raw_html)
    text = re.sub(r"(?i)<br\s*/?>", "\n", text)
    text = re.sub(r"(?i)</p>", "\n\n", text)
    text = re.sub(r"(?i)</div>", "\n", text)
    text = re.sub(r"(?i)</section>", "\n", text)
    text = re.sub(r"(?i)</h[1-6]>", "\n\n", text)
    text = re.sub(r"<[^>]+>", " ", text)
    text = html.unescape(text)
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
    return text.strip()


def wechat_crawler_path() -> Path:
    configured = os.environ.get("LLM_WIKI_WECHAT_CRAWLER", "").strip()
    return Path(configured).expanduser() if configured else DEFAULT_WECHAT_CRAWLER


def fetch_wechat_with_crawler(url: str, timeout: int = 45) -> tuple[str, str, dict[str, str]]:
    crawler_dir = wechat_crawler_path()
    if not crawler_dir.exists():
        return "", f"wechat crawler unavailable: {crawler_dir}", {}
    if not (crawler_dir / "enhanced_web_crawler.py").exists():
        return "", f"wechat crawler missing enhanced_web_crawler.py: {crawler_dir}", {}

    code = r"""
import json
import sys
from bs4 import BeautifulSoup
from enhanced_web_crawler import EnhancedWebCrawler

url = sys.argv[1]
crawler = EnhancedWebCrawler()
crawler.crawled_urls = set()
article = crawler.fetch_article(url)
if not article:
    print(json.dumps({"ok": False, "error": "crawler returned no article"}, ensure_ascii=False))
    raise SystemExit(0)

html_content = article.get("html_content", "")
soup = BeautifulSoup(html_content, "lxml")
text = soup.get_text("\n", strip=True)
print(json.dumps({
    "ok": True,
    "title": article.get("title", ""),
    "author": article.get("author", ""),
    "publish_time": article.get("publish_time", ""),
    "html_length": len(html_content),
    "text": text,
}, ensure_ascii=False))
"""
    try:
        completed = subprocess.run(
            [sys.executable, "-c", code, url],
            cwd=crawler_dir,
            check=False,
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=timeout,
        )
    except subprocess.TimeoutExpired:
        return "", "wechat crawler failed: timeout", {}
    except Exception as exc:  # noqa: BLE001
        return "", f"wechat crawler failed: {exc}", {}

    if completed.returncode != 0 and not completed.stdout.strip():
        return "", f"wechat crawler failed: {completed.stderr.strip()}", {}

    payload_line = ""
    for line in reversed(completed.stdout.splitlines()):
        if line.strip().startswith("{"):
            payload_line = line.strip()
            break
    if not payload_line:
        return "", f"wechat crawler failed: no JSON output; {completed.stderr.strip()}", {}

    try:
        payload = json.loads(payload_line)
    except json.JSONDecodeError:
        return "", f"wechat crawler failed: invalid JSON output; {payload_line[:120]}", {}

    if not payload.get("ok"):
        return "", f"wechat crawler failed: {payload.get('error', 'unknown error')}", {}

    text = str(payload.get("text", "")).strip()
    if is_blocked_or_boilerplate(text):
        return "", "blocked: wechat crawler fetched verification or anti-bot page", {}
    if not text:
        return "", "wechat crawler failed: empty article text", {}

    info = {
        "title": str(payload.get("title", "")).strip(),
        "author": str(payload.get("author", "")).strip(),
        "publish_time": str(payload.get("publish_time", "")).strip(),
        "html_length": str(payload.get("html_length", "")),
    }
    return text, "wechat crawler fetched", info


def fetch_web_text(url: str, timeout: int = 8) -> tuple[str, str]:
    if not url:
        return "", "skipped"
    try:
        request = urllib.request.Request(
            url,
            headers={"User-Agent": "LLM-Wiki-Auto-Ingest/0.1"},
        )
        with urllib.request.urlopen(request, timeout=timeout) as response:
            content_type = response.headers.get("content-type", "")
            raw = response.read(300_000)
        if "html" not in content_type and b"<html" not in raw[:1000].lower():
            return raw.decode("utf-8", errors="ignore"), "fetched non-html"
        text = html_to_text(raw.decode("utf-8", errors="ignore"))
        if is_blocked_or_boilerplate(text):
            return "", "blocked: fetched verification or anti-bot page"
        return text.strip(), "fetched"
    except Exception as exc:  # noqa: BLE001
        return "", f"fetch failed: {exc}"


def read_docx_text(path: Path) -> tuple[str, str]:
    try:
        from docx import Document
    except Exception as exc:  # noqa: BLE001
        return "", f"docx read failed: python-docx unavailable: {exc}"

    try:
        document = Document(path)
        blocks: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                blocks.append(text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))
        if not blocks:
            return "", "docx read: empty document"
        return "\n\n".join(blocks), "docx read"
    except Exception as exc:  # noqa: BLE001
        return "", f"docx read failed: {exc}"


def read_local_text(local_path: str) -> tuple[str, str]:
    if not local_path:
        return "", "skipped"
    path = Path(local_path).expanduser()
    if not path.exists():
        return "", "local path not found"
    if path.suffix.lower() == ".docx":
        return read_docx_text(path)
    if path.suffix.lower() not in {".md", ".markdown", ".txt", ".text", ".csv", ".tsv"}:
        return "", f"not extracted: unsupported file type {path.suffix or 'unknown'}"
    try:
        return path.read_text(encoding="utf-8", errors="ignore"), "read"
    except Exception as exc:  # noqa: BLE001
        return "", f"read failed: {exc}"


def update_source_status(path: Path, status: str) -> None:
    lines = path.read_text(encoding="utf-8").splitlines()
    if not lines or lines[0].strip() != "---":
        return

    updated: list[str] = [lines[0]]
    saw_status = False
    closed = False
    for line in lines[1:]:
        if line.strip() == "---" and not closed:
            if not saw_status:
                updated.append(f"status: {status}")
            updated.append(line)
            closed = True
            continue
        if not closed and line.startswith("status:"):
            updated.append(f"status: {status}")
            saw_status = True
            continue
        updated.append(line)
    path.write_text("\n".join(updated).rstrip() + "\n", encoding="utf-8")


def target_summary_path(raw_path: Path, created: str) -> Path:
    return WIKI_SOURCES / f"{created}-{slug_from_raw_path(raw_path)}.md"


def related_links(domain_slug: str) -> list[str]:
    links = []
    domain_page = WIKI_DOMAINS / f"{domain_slug}.md"
    if domain_page.exists():
        links.append(f"- [{domain_page.stem} Knowledge Network](../domains/{domain_slug}.md)")
    concept_page = WIKI_CONCEPTS / f"{domain_slug}.md"
    if concept_page.exists():
        links.append(f"- [{concept_page.stem}](../concepts/{domain_slug}.md)")
    return links


def build_summary(raw_path: Path, fetch_web: bool = True, read_local: bool = True) -> tuple[str, dict[str, str]]:
    meta, body = parse_frontmatter_and_body(raw_path)
    title = meta.get("title") or raw_path.stem
    kind = meta.get("kind") or "source"
    source_type = meta.get("source_type") or kind
    domain_slug = meta.get("domain") or raw_path.parent.parent.name
    domain_label = meta.get("domain_label") or domain_slug
    created = date_from_meta_or_path(meta, raw_path)
    today = dt.date.today().isoformat()
    tags = ["source", domain_slug]
    for tag in parse_list(meta.get("tags")):
        if tag not in tags:
            tags.append(tag)

    sections = extract_sections(body)
    context = remove_placeholders(sections.get("Context", ""))
    original = remove_placeholders(sections.get("Original Content Or Extract", ""))
    notes = remove_placeholders(sections.get("Notes", ""))

    if source_type == "sciverse":
        original = clean_sciverse_original(original)

    external_text = ""
    extraction_status = "skipped"
    crawler_info: dict[str, str] = {}
    if fetch_web and meta.get("source_url"):
        source_url = meta["source_url"]
        if "mp.weixin.qq.com" in source_url:
            external_text, extraction_status, crawler_info = fetch_wechat_with_crawler(source_url)
        if not external_text and not extraction_status.startswith("wechat crawler fetched"):
            fallback_text, fallback_status = fetch_web_text(source_url)
            if fallback_text:
                external_text = fallback_text
                extraction_status = fallback_status
            elif extraction_status == "skipped":
                extraction_status = fallback_status
    elif read_local and meta.get("local_path"):
        external_text, extraction_status = read_local_text(meta["local_path"])

    evidence_parts = [context, original, notes]
    if external_text and not is_blocked_or_boilerplate(external_text):
        evidence_parts.append(external_text)

    evidence_text = "\n\n".join(part for part in evidence_parts if part.strip())
    if evidence_text:
        summary = truncate(evidence_text, 180)
    elif extraction_status.startswith("blocked"):
        summary = f"已登记的{domain_label}网页资料，但正文抓取被验证页阻断；请在 raw source 的摘录区补充正文后重新 ingest。"
    elif extraction_status.startswith("fetch failed"):
        summary = f"已登记的{domain_label}网页资料，但正文抓取失败；请在 raw source 的摘录区补充正文后重新 ingest。"
    else:
        summary = f"已登记的{domain_label}资料，等待补充摘录或正文。"
    takeaways = first_paragraphs("\n\n".join(part for part in [context, original, notes] if part), 5)
    if len(takeaways) < 3 and external_text and not is_blocked_or_boilerplate(external_text):
        takeaways.extend(first_paragraphs(external_text, 5 - len(takeaways)))

    candidate_concepts = [
        tag
        for tag in tags
        if tag not in NOISE_CANDIDATE_TAGS and tag != domain_slug
    ]
    if not candidate_concepts and source_type != "sciverse":
        candidate_concepts = [domain_label, kind, source_type]

    raw_rel = raw_path.relative_to(ROOT).as_posix()
    locator_lines = [
        f"- Raw file: `{raw_rel}`",
        f"- Kind: `{kind}`",
        f"- Source type: `{source_type}`",
        f"- Domain: `{domain_slug}` / {domain_label}",
    ]
    if meta.get("source_url"):
        locator_lines.append(f"- URL: {meta['source_url']}")
    if meta.get("local_path"):
        locator_lines.append(f"- Local path: `{meta['local_path']}`")
    if meta.get("author"):
        locator_lines.append(f"- Author: {meta['author']}")
    elif crawler_info.get("author"):
        locator_lines.append(f"- Author: {crawler_info['author']}")
    if meta.get("published"):
        locator_lines.append(f"- Published: {meta['published']}")
    elif crawler_info.get("publish_time"):
        locator_lines.append(f"- Published: {crawler_info['publish_time']}")
    if crawler_info.get("html_length"):
        locator_lines.append(f"- WeChat crawler HTML length: {crawler_info['html_length']}")
    locator_lines.append(f"- Extraction: {extraction_status}")
    locator_lines.append("- Ingest mode: local-rule")

    source_excerpt_text = external_text if external_text and not is_blocked_or_boilerplate(external_text) else ""
    source_excerpt = first_paragraphs(source_excerpt_text or original or notes or context, 4)
    related = related_links(domain_slug)

    content_lines = [
        "---",
        f"title: Source - {title}",
        "type: source",
        "status: active",
        f"created: {today}",
        f"updated: {today}",
        f"summary: {summary}",
        f"tags: [{', '.join(tags)}]",
        f"sources: [{raw_rel}]",
        "---",
        "",
        f"# Source - {title}",
        "",
        "## Source Profile",
        "",
        *locator_lines,
        "",
        "## Why This Source Was Added",
        "",
        context or "- 待补：需要补充这份资料与当前 wiki 主线的关系。",
        "",
        "## Main Takeaways",
        "",
        *markdown_list(takeaways, "待补：需要人工或 LLM 继续阅读后提炼核心观点。"),
        "",
        "## Extracted Notes",
        "",
        *markdown_list(source_excerpt, "暂无可提取摘录。可以在 raw source 中补充 Original Content Or Extract 后重新 ingest。"),
        "",
        "## Candidate Concepts And Links",
        "",
        *[f"- `{concept}`" for concept in candidate_concepts],
        "",
        "## Next Integration Steps",
        "",
        "- 判断是否需要新建或更新 concept/entity 页面。",
        "- 把这份 source summary 链接到相关概念页的 Related 与 Sources 区块。",
        "- 若它暴露跨学科桥接点，把桥接关系写入相关 domain network 或 concept 页面。",
        "",
        "## Related",
        "",
        *(related or ["- 待补：对应 domain network 尚未生成。"]),
        "",
        "## Sources",
        "",
        f"- `{raw_rel}`",
    ]

    return "\n".join(content_lines).rstrip() + "\n", {
        "title": title,
        "domain": domain_slug,
        "summary": summary,
        "created": created,
        "raw_rel": raw_rel,
    }


def ingest(raw_source: Path, overwrite: bool, update_status: bool, fetch_web: bool, read_local: bool) -> dict[str, str]:
    raw_path = raw_source if raw_source.is_absolute() else ROOT / raw_source
    raw_path = raw_path.resolve()
    if not raw_path.exists():
        raise FileNotFoundError(raw_path)
    if RAW_SOURCES.resolve() not in raw_path.parents:
        raise ValueError(f"Source must be under {RAW_SOURCES.relative_to(ROOT)}")

    content, info = build_summary(raw_path, fetch_web=fetch_web, read_local=read_local)
    target = target_summary_path(raw_path, info["created"])
    if target.exists() and not overwrite:
        raise FileExistsError(f"File already exists: {target.relative_to(ROOT)}")

    WIKI_SOURCES.mkdir(parents=True, exist_ok=True)
    target.write_text(content, encoding="utf-8")
    if update_status:
        update_source_status(raw_path, "active")

    return {
        "raw_source": str(raw_path.relative_to(ROOT)),
        "summary": str(target.relative_to(ROOT)),
        "domain": info["domain"],
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate a wiki source summary from a raw source.")
    parser.add_argument("source", help="Raw source markdown path.")
    parser.add_argument("--overwrite", action="store_true", help="Overwrite existing source summary.")
    parser.add_argument("--keep-status", action="store_true", help="Do not mark raw source as active.")
    parser.add_argument("--no-fetch-web", action="store_true", help="Do not fetch source_url content.")
    parser.add_argument("--no-read-local", action="store_true", help="Do not read text local_path content.")
    args = parser.parse_args()

    result = ingest(
        Path(args.source),
        overwrite=args.overwrite,
        update_status=not args.keep_status,
        fetch_web=not args.no_fetch_web,
        read_local=not args.no_read_local,
    )
    print(result["summary"])
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:  # noqa: BLE001
        print(f"ERROR: {exc}", file=sys.stderr)
        raise SystemExit(1)
